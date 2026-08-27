"""Run a safe, fictional end-to-end latency benchmark for /api/analyze."""

import logging
import re
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


def fictional_resume() -> bytes:
    document = Document()
    document.add_heading("Alex Example", level=1)
    document.add_paragraph("AI Product Manager | Sydney, Australia")
    document.add_heading("Experience", level=2)
    document.add_paragraph("Senior Product Manager — Example Labs | 2023–2026")
    document.add_paragraph("Led discovery for an AI support assistant with 500 fictional users.")
    document.add_paragraph(
        "Partnered with engineering, design, and data teams from pilot to launch."
    )
    document.add_paragraph("Product Manager — Sample Cloud | 2021–2023")
    document.add_paragraph("Owned roadmap prioritization for a B2B analytics product.")
    document.add_paragraph(
        "Used SQL dashboards and customer interviews to guide quarterly planning."
    )
    document.add_heading("Projects", level=2)
    document.add_paragraph(
        "JobPilot — Built a fictional resume-to-role matching prototype using Python."
    )
    document.add_heading("Education and skills", level=2)
    document.add_paragraph("BSc Computer Science — Example University | 2021")
    document.add_paragraph("Product strategy, customer discovery, AI, SQL, Python, analytics")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def main() -> None:
    logging.basicConfig(level=logging.WARNING, format="%(levelname)s %(name)s %(message)s")
    logging.getLogger("app").setLevel(logging.INFO)
    response = TestClient(app).post(
        "/api/analyze",
        files={
            "resume": (
                "fictional-resume.docx",
                fictional_resume(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "target_position": "AI Product Manager",
            "job_description": (
                "Example Labs is hiring an AI Product Manager to own product strategy for an "
                "enterprise AI assistant. Responsibilities include customer discovery, defining "
                "product requirements, roadmap prioritization, experiment design, launch planning, "
                "and cross-functional delivery with engineering, design, data science, sales, and "
                "legal. Required qualifications include three years of product management, evidence "
                "of shipping AI-powered products, strong analytical judgment, executive communication, "
                "and the ability to translate customer problems into measurable outcomes. Experience "
                "with LLM evaluation, responsible AI, SQL, Python, and B2B SaaS is preferred."
            ),
        },
    )
    body = response.json()
    print(f"benchmark_status={response.status_code}")
    if response.status_code != 200:
        print(f"benchmark_error_type={type(body.get('detail')).__name__}")
        raise SystemExit(1)
    match = body["match_analysis"]
    visible_analysis = " ".join(
        match["top_strengths"]
        + match["key_gaps"]
        + match["suggested_preparation"]
        + [item["requirement"] for item in match["evidence"]]
        + [item["resume_evidence"] for item in match["evidence"]]
    )
    chinese_characters = len(re.findall("[一-龥]", visible_analysis))
    print("benchmark_structured_output=true")
    print(f"benchmark_chinese_characters={chinese_characters}")
    print(f"benchmark_recommendation_enum={match['recommendation']}")


if __name__ == "__main__":
    main()
