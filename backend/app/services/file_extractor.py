from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class FileExtractionError(ValueError):
    pass


def extract_resume_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            document = Document(BytesIO(content))
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                parts.extend(cell.text for row in table.rows for cell in row.cells)
            text = "\n".join(parts)
        else:
            raise FileExtractionError("Only PDF and DOCX resumes are supported.")
    except FileExtractionError:
        raise
    except Exception as exc:
        raise FileExtractionError(
            "The resume could not be read. Check that the file is valid."
        ) from exc

    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not cleaned:
        raise FileExtractionError(
            "No text was found in the resume. Scanned PDFs are not supported in V0.1."
        )
    return cleaned
