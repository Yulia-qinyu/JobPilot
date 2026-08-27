from io import BytesIO

import pytest
from docx import Document

from app.services.file_extractor import FileExtractionError, extract_resume_text


def test_extracts_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Ada Lovelace")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "AI Product Manager"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_resume_text("resume.docx", buffer.getvalue())

    assert "Ada Lovelace" in text
    assert "AI Product Manager" in text


def test_rejects_unsupported_file_type() -> None:
    with pytest.raises(FileExtractionError, match="Only PDF and DOCX"):
        extract_resume_text("resume.txt", b"hello")
